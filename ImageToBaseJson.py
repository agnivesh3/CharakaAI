import os
import re
import json
import logging
from typing import List, Dict, Any
from datetime import datetime
from google.cloud import documentai_v1 as documentai
from docx import Document

# Set up logging
log_folder = "logs"
os.makedirs(log_folder, exist_ok=True)
log_file = os.path.join(log_folder, f"charaka_samhita_processing_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log")
logging.basicConfig(filename=log_file, level=logging.DEBUG,
                    format='%(asctime)s - %(levelname)s - %(message)s')

def process_document(project_id: str, location: str, processor_id: str, file_path: str) -> documentai.Document:
    logging.info(f"Processing document: {file_path}")
    client = documentai.DocumentProcessorServiceClient()
    name = f"projects/{project_id}/locations/{location}/processors/{processor_id}"

    with open(file_path, "rb") as image:
        image_content = image.read()

    raw_document = documentai.RawDocument(content=image_content, mime_type="image/png")
    request = documentai.ProcessRequest(name=name, raw_document=raw_document)

    try:
        result = client.process_document(request=request)
        logging.info("Document processed successfully")
        return result.document
    except Exception as e:
        logging.error(f"Error processing document: {e}")
        raise

def extract_text(document: documentai.Document) -> str:
    """Extract text from the entire document."""
    return document.text

def is_sanskrit(text: str) -> bool:
    return bool(re.search(r'[\u0900-\u097F]', text))

def extract_verses_and_groups(ocr_text: str) -> tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
    sanskrit_verses = []
    verse_groups = []
    current_verse = ""
    current_verse_number = None

    lines = ocr_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        # Detect verse groups
        verse_group_match = re.search(r'\[(\d+(?:-\d+)?)\]', line)
        if verse_group_match:
            verse_groups.append(verse_group_match.group(1))

        # Check for verse end
        verse_end_match = re.search(r'।।\s*(\d+)\s*।।', line)
        if verse_end_match:
            verse_number = int(verse_end_match.group(1))
            current_verse += " " + line[:verse_end_match.start()].strip()
            if is_sanskrit(current_verse):
                sanskrit_verses.append({
                    "verse_number": verse_number,
                    "sanskrit": clean_sanskrit(current_verse)
                })
            current_verse = ""
            current_verse_number = None
        elif current_verse_number is not None:
            current_verse += " " + line
        else:
            # Check for verse start
            verse_start_match = re.match(r'^(\d+)\.?\s*(.+)', line)
            if verse_start_match:
                if current_verse and is_sanskrit(current_verse):
                    sanskrit_verses.append({
                        "verse_number": current_verse_number,
                        "sanskrit": clean_sanskrit(current_verse)
                    })
                current_verse_number = int(verse_start_match.group(1))
                current_verse = verse_start_match.group(2)
            elif is_sanskrit(line):
                current_verse = line
                current_verse_number = len(sanskrit_verses) + 1

        i += 1

    # Handle any remaining verse
    if current_verse and is_sanskrit(current_verse):
        sanskrit_verses.append({
            "verse_number": current_verse_number or len(sanskrit_verses) + 1,
            "sanskrit": clean_sanskrit(current_verse)
        })

    return sanskrit_verses, verse_groups

def clean_sanskrit(text: str) -> str:
    # Remove non-Devanagari characters except for । and ॥
    cleaned = re.sub(r'[^\u0900-\u097F\s।॥]', '', text)
    # Remove extra spaces
    cleaned = re.sub(r'\s+', ' ', cleaned).strip()
    return cleaned

def process_all_images(project_id: str, location: str, processor_id: str, folder_path: str) -> tuple[List[Dict[str, Any]], List[Dict[str, Any]], List[str]]:
    all_ocr_text = ""
    image_files = sorted([f for f in os.listdir(folder_path) if f.endswith('.png')])

    for image_file in image_files:
        file_path = os.path.join(folder_path, image_file)
        document = process_document(project_id, location, processor_id, file_path)
        all_ocr_text += extract_text(document) + "\n\n"

    sanskrit_verses, verse_groups = extract_verses_and_groups(all_ocr_text)
    return sanskrit_verses, verse_groups, [all_ocr_text]

def save_to_word(sanskrit_verses: List[Dict[str, Any]], verse_groups: List[Dict[str, Any]], ocr_content: List[str], output_path: str) -> None:
    document = Document()
    document.add_heading('Charaka Samhita OCR Output', 0)

    document.add_heading('OCR Content', level=1)
    for content in ocr_content:
        document.add_paragraph(content)

    document.add_heading('Sanskrit Verses', level=1)
    for verse in sanskrit_verses:
        if verse.get('error'):
            document.add_paragraph(f"Error: {verse['error']}")
        document.add_paragraph(f"Verse {verse.get('verse_number', 'Unknown')}: {verse['sanskrit']}")

    document.add_heading('Verse Groups', level=1)
    for group in verse_groups:
        document.add_paragraph(f"Verse Group: {group}")

    document.save(output_path)
    logging.info(f"OCR output has been written to {output_path}")

def process_text(ocr_text: str) -> dict:
    sanskrit_verses, verse_groups = extract_verses_and_groups(ocr_text)

    output = {
        "timestamp": datetime.now().isoformat(),
        "book": {
            "title": "Charaka Samhita",
            "volume": 1,
            "section": 1,
            "chapter": 5,
            "chapter_name": "",
            "sanskrit_verses": sanskrit_verses,
            "verse_groups": verse_groups
        }
    }

    return output

# Main execution
if __name__ == "__main__":
    # Set Google Cloud credentials
    os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = "csprocessor-service-account.json"

    # Google Cloud project details
    project_id = "1039812532642"
    location = "us"
    processor_id = "1fb139ed60959696"
    folder_path = "ExtractedImage/S1-Chapter5"

    try:
        # Process all images in the folder
        sanskrit_verses, verse_groups, ocr_content = process_all_images(project_id, location, processor_id, folder_path)

        # Log the extracted content for debugging
        logging.debug(f"Extracted Sanskrit verses: {sanskrit_verses}")
        logging.debug(f"Extracted verse groups: {verse_groups}")

        # Process the OCR text
        if ocr_content and ocr_content[0]:
            result = process_text(ocr_content[0])
        else:
            logging.error("No OCR content extracted")
            raise ValueError("No OCR content extracted")

        # Create output folder structure
        section = result['book']['section']
        chapter = result['book']['chapter']
        output_folder = f"ExtractedFromOCR/S{section}C{chapter}"
        os.makedirs(output_folder, exist_ok=True)

        # Write the JSON output to a file
        json_output_path = os.path.join(output_folder, 'charaka_samhita_output.json')
        with open(json_output_path, 'w', encoding='utf-8') as f:
            json.dump(result, f, ensure_ascii=False, indent=2)

        logging.info(f"JSON output has been written to {json_output_path}")

        # Save the OCR output to a Word document for readability
        docx_output_path = os.path.join(output_folder, 'charaka_samhita_output.docx')
        save_to_word(sanskrit_verses, verse_groups, ocr_content, docx_output_path)

        logging.info(f"Word document has been saved to {docx_output_path}")

    except Exception as e:
        logging.error(f"An error occurred: {e}")
        raise
